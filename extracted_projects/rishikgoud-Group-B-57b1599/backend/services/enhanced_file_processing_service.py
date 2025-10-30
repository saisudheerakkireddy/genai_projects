"""
Enhanced File Processing Service
Industry-standard PDF and document processing with comprehensive validation
"""

import os
import logging
import hashlib
from typing import Dict, Any, List, Optional
from datetime import datetime
import asyncio
import json

# PDF processing libraries
try:
    import PyPDF2
    from PyPDF2 import PdfReader
except ImportError:
    PyPDF2 = None
    PdfReader = None

try:
    import pdfplumber
except ImportError:
    pdfplumber = None

try:
    from pdf2image import convert_from_path
except ImportError:
    convert_from_path = None

try:
    import pytesseract
    from PIL import Image
except ImportError:
    pytesseract = None
    Image = None

try:
    from docx import Document
except ImportError:
    Document = None

logger = logging.getLogger(__name__)

class EnhancedFileProcessingService:
    """Comprehensive file processing service with industry-standard features"""

    def __init__(self):
        """Initialize enhanced file processing service"""

        # Configure Tesseract if available
        if pytesseract and os.name == 'nt':
            # Set Tesseract path for Windows
            try:
                pytesseract.pytesseract.tesseract_cmd = r'C:\Program Files\Tesseract-OCR\tesseract.exe'
            except:
                pass

        # File processing limits
        self.max_file_size = 50 * 1024 * 1024  # 50MB
        self.max_pages = 500
        self.supported_formats = {
            'pdf': ['.pdf'],
            'docx': ['.docx', '.doc'],
            'txt': ['.txt'],
            'images': ['.png', '.jpg', '.jpeg', '.tiff', '.bmp']
        }

    async def process_uploaded_contract(
        self,
        file_path: str,
        filename: str,
        file_size: int,
        file_type: str,
        contract_id: str
    ) -> Dict[str, Any]:
        """
        Comprehensive contract processing pipeline

        Args:
            file_path: Path to uploaded file
            filename: Original filename
            file_size: File size in bytes
            file_type: File extension
            contract_id: Contract database ID

        Returns:
            Complete processing results with metadata
        """

        logger.info(f"Starting comprehensive processing for {filename}")

        processing_result = {
            "contract_id": contract_id,
            "filename": filename,
            "file_size": file_size,
            "file_type": file_type,
            "processing_status": "processing",
            "start_time": datetime.utcnow(),
            "steps": [],
            "metadata": {},
            "extracted_text": "",
            "quality_metrics": {},
            "errors": [],
            "warnings": []
        }

        try:
            # Step 1: File validation
            validation_result = await self._validate_file_comprehensive(file_path, file_type)
            processing_result["validation"] = validation_result
            processing_result["steps"].append("file_validation")

            if not validation_result["valid"]:
                processing_result["processing_status"] = "failed"
                processing_result["errors"] = validation_result["errors"]
                return processing_result

            # Step 2: Security scanning
            security_result = await self._security_scan(file_path)
            processing_result["security"] = security_result
            processing_result["steps"].append("security_scan")

            # Step 3: Text extraction
            extraction_result = await self._extract_text_enhanced(file_path, file_type)
            processing_result["extraction"] = extraction_result
            processing_result["extracted_text"] = extraction_result["text"]
            processing_result["steps"].append("text_extraction")

            # Step 4: Document analysis
            document_analysis = await self._analyze_document_structure(extraction_result["text"])
            processing_result["document_analysis"] = document_analysis
            processing_result["steps"].append("document_analysis")

            # Step 5: Quality assessment
            quality_metrics = self._assess_quality(extraction_result, document_analysis)
            processing_result["quality_metrics"] = quality_metrics
            processing_result["steps"].append("quality_assessment")

            # Step 6: Metadata extraction
            metadata = await self._extract_metadata(file_path, file_type, extraction_result)
            processing_result["metadata"] = metadata
            processing_result["steps"].append("metadata_extraction")

            # Final status
            processing_result["processing_status"] = "completed"
            processing_result["end_time"] = datetime.utcnow()
            processing_result["processing_time"] = (
                processing_result["end_time"] - processing_result["start_time"]
            ).total_seconds()

            logger.info(f"Processing completed for {filename}")
            return processing_result

        except Exception as e:
            logger.error(f"Processing failed for {filename}: {e}")
            processing_result["processing_status"] = "failed"
            processing_result["errors"].append(str(e))
            processing_result["end_time"] = datetime.utcnow()
            return processing_result

    async def _validate_file_comprehensive(self, file_path: str, file_type: str) -> Dict[str, Any]:
        """Comprehensive file validation"""

        validation = {
            "valid": True,
            "errors": [],
            "warnings": [],
            "file_info": {}
        }

        try:
            # Basic file existence check
            if not os.path.exists(file_path):
                validation["valid"] = False
                validation["errors"].append("File does not exist")
                return validation

            # File size validation
            actual_size = os.path.getsize(file_path)
            if actual_size > self.max_file_size:
                validation["valid"] = False
                validation["errors"].append(f"File too large: {actual_size} bytes (max: {self.max_file_size})")

            # File type validation
            if file_type.lower() not in [ft for sublist in self.supported_formats.values() for ft in sublist]:
                validation["valid"] = False
                validation["errors"].append(f"Unsupported file type: {file_type}")

            # PDF-specific validation
            if file_type.lower() == 'pdf':
                validation.update(await self._validate_pdf(file_path))

            # Image-specific validation
            elif file_type.lower() in ['png', 'jpg', 'jpeg']:
                validation.update(await self._validate_image(file_path))

            # Calculate file checksum
            validation["file_info"]["checksum"] = await self._calculate_checksum(file_path)

        except Exception as e:
            validation["valid"] = False
            validation["errors"].append(f"Validation error: {str(e)}")

        return validation

    async def _validate_pdf(self, file_path: str) -> Dict[str, Any]:
        """Validate PDF file structure"""

        validation = {"pdf_info": {}}

        try:
            if pdfplumber:
                with pdfplumber.open(file_path) as pdf:
                    page_count = len(pdf.pages)
                    validation["pdf_info"]["page_count"] = page_count

                    if page_count > self.max_pages:
                        validation.setdefault("warnings", []).append(
                            f"Large document: {page_count} pages (max recommended: {self.max_pages})"
                        )

                    # Check for text content
                    has_text = False
                    for page in pdf.pages:
                        if page.extract_text():
                            has_text = True
                            break

                    if not has_text:
                        validation.setdefault("warnings", []).append(
                            "PDF appears to contain only images (OCR may be required)"
                        )

            elif PyPDF2:
                with open(file_path, 'rb') as file:
                    pdf = PyPDF2.PdfReader(file)
                    validation["pdf_info"]["page_count"] = len(pdf.pages)

        except Exception as e:
            validation.setdefault("errors", []).append(f"PDF validation failed: {str(e)}")

        return validation

    async def _validate_image(self, file_path: str) -> Dict[str, Any]:
        """Validate image file"""

        validation = {"image_info": {}}

        try:
            if Image:
                with Image.open(file_path) as img:
                    validation["image_info"]["size"] = f"{img.width}x{img.height}"
                    validation["image_info"]["mode"] = img.mode
                    validation["image_info"]["format"] = img.format

                    # Check image quality
                    if img.width < 100 or img.height < 100:
                        validation.setdefault("warnings", []).append(
                            "Low resolution image may result in poor OCR quality"
                        )

        except Exception as e:
            validation.setdefault("errors", []).append(f"Image validation failed: {str(e)}")

        return validation

    async def _security_scan(self, file_path: str) -> Dict[str, Any]:
        """Basic security scanning for uploaded files"""

        security_result = {
            "scan_status": "completed",
            "threats_found": 0,
            "scan_details": {}
        }

        # Note: In production, integrate with actual antivirus software
        # For now, basic checks

        try:
            # Check file size (unusually large files)
            file_size = os.path.getsize(file_path)
            if file_size > 100 * 1024 * 1024:  # 100MB
                security_result.setdefault("warnings", []).append("Unusually large file size")

            # Check for suspicious file extensions
            filename = os.path.basename(file_path).lower()
            suspicious_patterns = ['.exe', '.bat', '.scr', '.vbs', '.js']
            if any(pattern in filename for pattern in suspicious_patterns):
                security_result["threats_found"] += 1
                security_result.setdefault("threats", []).append("Suspicious file extension")

        except Exception as e:
            security_result["scan_status"] = "failed"
            security_result["error"] = str(e)

        return security_result

    async def _extract_text_enhanced(self, file_path: str, file_type: str) -> Dict[str, Any]:
        """Enhanced text extraction with multiple fallback methods"""

        extraction_result = {
            "text": "",
            "extraction_method": "unknown",
            "success": False,
            "pages_info": [],
            "quality_score": 0.0
        }

        try:
            if file_type.lower() == 'pdf':
                extraction_result = await self._extract_from_pdf_enhanced(file_path)
            elif file_type.lower() in ['png', 'jpg', 'jpeg']:
                extraction_result = await self._extract_from_image_enhanced(file_path)
            elif file_type.lower() == 'txt':
                extraction_result = await self._extract_from_text_enhanced(file_path)
            elif file_type.lower() == 'docx':
                extraction_result = await self._extract_from_docx_enhanced(file_path)
            else:
                raise ValueError(f"Unsupported file type: {file_type}")

            # Calculate quality score
            if extraction_result["text"]:
                extraction_result["quality_score"] = self._calculate_extraction_quality(
                    extraction_result["text"]
                )
                extraction_result["success"] = True

        except Exception as e:
            logger.error(f"Enhanced text extraction failed: {e}")
            extraction_result["error"] = str(e)

        return extraction_result

    async def _extract_from_pdf_enhanced(self, file_path: str) -> Dict[str, Any]:
        """Enhanced PDF text extraction with multiple methods and quality assessment"""

        text_content = ""
        pages_info = []
        methods_tried = []

        # Method 1: pdfplumber (best for complex layouts)
        if pdfplumber:
            try:
                with pdfplumber.open(file_path) as pdf:
                    for page_num, page in enumerate(pdf.pages, 1):
                        page_text = page.extract_text()
                        if page_text and page_text.strip():
                            text_content += f"\n--- Page {page_num} ---\n{page_text}\n"
                            pages_info.append({
                                "page_number": page_num,
                                "extraction_method": "pdfplumber",
                                "text_length": len(page_text),
                                "has_text": True,
                                "quality": self._assess_page_quality(page_text)
                            })
                        else:
                            pages_info.append({
                                "page_number": page_num,
                                "extraction_method": "pdfplumber",
                                "text_length": 0,
                                "has_text": False,
                                "quality": 0.0
                            })

                if text_content.strip():
                    methods_tried.append("pdfplumber")
                    return {
                        "text": text_content.strip(),
                        "extraction_method": "pdfplumber",
                        "pages_info": pages_info,
                        "total_pages": len(pages_info),
                        "methods_tried": methods_tried,
                        "success": True
                    }
            except Exception as e:
                logger.warning(f"pdfplumber failed: {e}")

        # Method 2: PyPDF2 (fallback)
        if PyPDF2 and not text_content.strip():
            try:
                with open(file_path, 'rb') as file:
                    pdf_reader = PdfReader(file)
                    for page_num, page in enumerate(pdf_reader.pages, 1):
                        page_text = page.extract_text()
                        if page_text and page_text.strip():
                            text_content += f"\n--- Page {page_num} ---\n{page_text}\n"
                            pages_info.append({
                                "page_number": page_num,
                                "extraction_method": "pypdf2",
                                "text_length": len(page_text),
                                "has_text": True,
                                "quality": self._assess_page_quality(page_text)
                            })

                if text_content.strip():
                    methods_tried.append("pypdf2")
                    return {
                        "text": text_content.strip(),
                        "extraction_method": "pypdf2",
                        "pages_info": pages_info,
                        "total_pages": len(pages_info),
                        "methods_tried": methods_tried,
                        "success": True
                    }
            except Exception as e:
                logger.warning(f"PyPDF2 failed: {e}")

        # Method 3: OCR (if images are present)
        if convert_from_path and pytesseract and not text_content.strip():
            try:
                pages = convert_from_path(file_path, dpi=300)
                for page_num, page in enumerate(pages, 1):
                    page_text = pytesseract.image_to_string(page, lang='eng')
                    if page_text and page_text.strip():
                        text_content += f"\n--- Page {page_num} (OCR) ---\n{page_text}\n"
                        pages_info.append({
                            "page_number": page_num,
                            "extraction_method": "ocr",
                            "text_length": len(page_text),
                            "has_text": True,
                            "quality": self._assess_page_quality(page_text)
                        })

                if text_content.strip():
                    methods_tried.append("ocr")
                    return {
                        "text": text_content.strip(),
                        "extraction_method": "ocr",
                        "pages_info": pages_info,
                        "total_pages": len(pages_info),
                        "methods_tried": methods_tried,
                        "success": True
                    }
            except Exception as e:
                logger.warning(f"OCR extraction failed: {e}")

        # If all methods failed
        return {
            "text": text_content,
            "extraction_method": "failed",
            "pages_info": pages_info,
            "total_pages": len(pages_info) if pages_info else 0,
            "methods_tried": methods_tried,
            "success": bool(text_content.strip()),
            "error": "All text extraction methods failed"
        }

    async def _extract_from_text_enhanced(self, file_path: str) -> Dict[str, Any]:
        """Enhanced text file extraction"""

        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                content = file.read()

            return {
                "text": content,
                "extraction_method": "direct_read",
                "success": True,
                "encoding": "utf-8",
                "line_count": len(content.split('\n'))
            }

        except UnicodeDecodeError:
            # Try different encodings
            encodings = ['latin-1', 'cp1252', 'iso-8859-1']
            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding) as file:
                        content = file.read()
                    return {
                        "text": content,
                        "extraction_method": f"direct_read_{encoding}",
                        "success": True,
                        "encoding": encoding,
                        "line_count": len(content.split('\n'))
                    }
                except:
                    continue

            raise Exception("Could not decode text file with any supported encoding")

    async def _extract_from_docx_enhanced(self, file_path: str) -> Dict[str, Any]:
        """Enhanced DOCX text extraction"""

        if not Document:
            raise Exception("python-docx library not available")

        try:
            doc = Document(file_path)
            content = []

            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    content.append(paragraph.text)

            # Extract tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        content.append(" | ".join(row_text))

            full_text = "\n".join(content)

            return {
                "text": full_text,
                "extraction_method": "docx",
                "success": True,
                "paragraph_count": len([p for p in doc.paragraphs if p.text.strip()]),
                "table_count": len(doc.tables)
            }

        except Exception as e:
            raise Exception(f"DOCX extraction failed: {str(e)}")

    async def _analyze_document_structure(self, text: str) -> Dict[str, Any]:
        """Analyze document structure and content"""

        analysis = {
            "word_count": len(text.split()),
            "character_count": len(text),
            "line_count": len(text.split('\n')),
            "paragraph_count": len([p for p in text.split('\n\n') if p.strip()]),
            "estimated_reading_time": len(text.split()) / 200,  # Average reading speed
        }

        # Detect potential sections
        sections = []
        lines = text.split('\n')
        current_section = None

        for line in lines:
            line = line.strip()
            if line and len(line) < 100 and line.isupper():
                if current_section:
                    sections.append(current_section)
                current_section = {"title": line, "content": []}
            elif current_section and line:
                current_section["content"].append(line)

        if current_section:
            sections.append(current_section)

        analysis["detected_sections"] = sections
        analysis["section_count"] = len(sections)

        return analysis

    def _calculate_extraction_quality(self, text: str) -> float:
        """Calculate text extraction quality score"""

        if not text:
            return 0.0

        score = 1.0

        # Check for OCR artifacts
        artifact_chars = ['|', '_', '*', '#', '@']
        artifact_count = sum(text.count(char) for char in artifact_chars)
        if artifact_count > 0:
            score -= min(0.3, artifact_count / len(text) * 10)

        # Check for reasonable sentence structure
        sentences = [s.strip() for s in text.split('.') if s.strip()]
        if len(sentences) < 2:
            score -= 0.2

        # Check for legal document patterns
        legal_indicators = [
            'agreement', 'contract', 'party', 'parties', 'shall', 'hereby',
            'whereas', 'therefore', 'pursuant', 'section', 'article'
        ]
        legal_count = sum(1 for indicator in legal_indicators if indicator.lower() in text.lower())
        if legal_count == 0:
            score -= 0.1

        return max(0.0, min(1.0, score))

    def _assess_page_quality(self, page_text: str) -> float:
        """Assess quality of individual page text"""

        if not page_text:
            return 0.0

        score = 1.0

        # Check text length
        if len(page_text) < 50:
            score -= 0.3
        elif len(page_text) > 5000:
            score -= 0.1

        # Check for meaningful words
        words = page_text.split()
        if len(words) < 10:
            score -= 0.2

        # Check for OCR artifacts
        artifacts = sum(page_text.count(char) for char in ['|', '_', '*'])
        if artifacts > 10:
            score -= 0.2

        return max(0.0, min(1.0, score))

    async def _calculate_checksum(self, file_path: str) -> str:
        """Calculate file checksum for integrity verification"""

        try:
            hash_md5 = hashlib.md5()
            with open(file_path, "rb") as f:
                for chunk in iter(lambda: f.read(4096), b""):
                    hash_md5.update(chunk)
            return hash_md5.hexdigest()
        except Exception as e:
            logger.error(f"Checksum calculation failed: {e}")
            return ""

    def _assess_quality(self, extraction_result: Dict[str, Any], document_analysis: Dict[str, Any]) -> Dict[str, Any]:
        """Assess overall processing quality"""

        quality = {
            "overall_score": 0.0,
            "extraction_quality": extraction_result.get("quality_score", 0.0),
            "document_structure": 0.0,
            "completeness": 0.0
        }

        # Document structure quality
        if document_analysis.get("section_count", 0) > 0:
            quality["document_structure"] = 0.8
        else:
            quality["document_structure"] = 0.4

        # Completeness assessment
        required_fields = ["text", "extraction_method", "pages_info"]
        completeness = sum(1 for field in required_fields if extraction_result.get(field))
        quality["completeness"] = completeness / len(required_fields)

        # Overall score
        quality["overall_score"] = (
            quality["extraction_quality"] * 0.5 +
            quality["document_structure"] * 0.3 +
            quality["completeness"] * 0.2
        )

        return quality

    async def _extract_metadata(self, file_path: str, file_type: str, extraction_result: Dict[str, Any]) -> Dict[str, Any]:
        """Extract comprehensive document metadata"""

        metadata = {
            "file_system": {},
            "document_properties": {},
            "processing_info": {}
        }

        try:
            # File system metadata
            stat = os.stat(file_path)
            metadata["file_system"] = {
                "created_time": datetime.fromtimestamp(stat.st_ctime),
                "modified_time": datetime.fromtimestamp(stat.st_mtime),
                "accessed_time": datetime.fromtimestamp(stat.st_atime),
                "permissions": oct(stat.st_mode)[-3:]
            }

            # Document-specific metadata
            if file_type.lower() == 'pdf':
                metadata["document_properties"] = await self._extract_pdf_metadata(file_path)
            elif file_type.lower() == 'docx':
                metadata["document_properties"] = await self._extract_docx_metadata(file_path)

            # Processing metadata
            metadata["processing_info"] = {
                "extraction_method": extraction_result.get("extraction_method"),
                "processing_timestamp": datetime.utcnow(),
                "text_quality_score": extraction_result.get("quality_score", 0.0),
                "total_pages": extraction_result.get("total_pages", 0)
            }

        except Exception as e:
            logger.error(f"Metadata extraction failed: {e}")
            metadata["error"] = str(e)

        return metadata

    async def _extract_pdf_metadata(self, file_path: str) -> Dict[str, Any]:
        """Extract PDF-specific metadata"""

        metadata = {}

        try:
            if pdfplumber:
                with pdfplumber.open(file_path) as pdf:
                    metadata["page_count"] = len(pdf.pages)
                    metadata["has_metadata"] = bool(pdf.metadata)

                    if pdf.metadata:
                        metadata["pdf_metadata"] = pdf.metadata

        except Exception as e:
            metadata["error"] = str(e)

        return metadata

    async def _extract_docx_metadata(self, file_path: str) -> Dict[str, Any]:
        """Extract DOCX-specific metadata"""

        if not Document:
            return {"error": "python-docx not available"}

        metadata = {}

        try:
            doc = Document(file_path)
            metadata["paragraph_count"] = len(doc.paragraphs)
            metadata["table_count"] = len(doc.tables)

            # Extract document properties if available
            if hasattr(doc.core_properties, 'author'):
                metadata["author"] = doc.core_properties.author
            if hasattr(doc.core_properties, 'created'):
                metadata["created_date"] = doc.core_properties.created
            if hasattr(doc.core_properties, 'modified'):
                metadata["modified_date"] = doc.core_properties.modified

        except Exception as e:
            metadata["error"] = str(e)

        return metadata
